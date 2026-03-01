from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Body
from sqlalchemy.orm import Session
from typing import List, Optional
from datetime import datetime, timedelta
import os
import shutil
import re
import openai
from PyPDF2 import PdfReader
from docx import Document

from models import Contract, ContractStatus, ContractType, ContractVersion, ReviewRule, Conversation, Message
from config import settings
from database import get_db

router = APIRouter()

# 配置OpenAI
if settings.OPENAI_API_KEY:
    openai.api_key = settings.OPENAI_API_KEY

# 配置DashScope API
import requests
import json

# 文件上传目录
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# 提取文件内容的函数
def extract_text_from_file(file_path, file_type):
    """从不同类型的文件中提取文本内容"""
    try:
        # 规范化文件路径，处理反斜杠问题
        file_path = os.path.normpath(file_path)
        print(f"尝试提取文件内容: {file_path}")
        print(f"文件是否存在: {os.path.exists(file_path)}")
        
        if not os.path.exists(file_path):
            return f"文件不存在: {file_path}"
        
        if file_type == "text/plain" or file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        elif file_type == "application/pdf" or file_path.endswith(".pdf"):
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"] or file_path.endswith(".docx"):
            try:
                doc = Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except Exception as e:
                print(f"读取Word文件失败: {str(e)}")
                # 如果Word文件读取失败，尝试将其作为文本文件读取
                try:
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        return f.read()
                except Exception as e2:
                    print(f"读取文本文件失败: {str(e2)}")
                    return f"读取文件失败: {str(e2)}"
        elif file_type == "application/msword" or file_path.endswith(".doc"):
            return "不支持的Word文件格式: .doc"
        else:
            return f"不支持的文件类型: {file_type}"
    except Exception as e:
        print(f"提取文件内容失败: {str(e)}")
        return f"提取文件内容失败: {str(e)}"

# 合同相关路由

# 获取合同列表
@router.get("/contracts", response_model=List[dict])
async def get_contracts(
    status: Optional[ContractStatus] = None,
    contract_type: Optional[ContractType] = None,
    start_date: Optional[datetime] = None,
    end_date: Optional[datetime] = None,
    db: Session = Depends(get_db)
):
    query = db.query(Contract)
    
    # 应用筛选条件
    if status:
        query = query.filter(Contract.status == status)
    if contract_type:
        query = query.filter(Contract.contract_type == contract_type)
    if start_date:
        query = query.filter(Contract.created_at >= start_date)
    if end_date:
        query = query.filter(Contract.created_at <= end_date)
    
    contracts = query.all()
    
    # 转换为响应格式
    result = []
    for contract in contracts:
        latest_version = contract.versions[0] if contract.versions else None
        result.append({
            "id": contract.id,
            "title": contract.title,
            "description": contract.description,
            "status": contract.status.value,
            "contract_type": contract.contract_type.value,
            "created_at": contract.created_at,
            "updated_at": contract.updated_at,
            "latest_version": latest_version.version if latest_version else None
        })
    
    return result

# 创建合同
@router.post("/contracts", response_model=dict)
async def create_contract(
    title: str = Form(...),
    description: Optional[str] = Form(None),
    contract_type: ContractType = Form(ContractType.OTHER),
    file: Optional[UploadFile] = File(None),
    db: Session = Depends(get_db)
):
    # 创建合同
    contract = Contract(
        title=title,
        description=description,
        contract_type=contract_type
    )
    db.add(contract)
    db.commit()
    db.refresh(contract)
    
    # 处理文件上传
    if file:
        # 保存文件
        file_path = os.path.join(UPLOAD_DIR, f"{contract.id}_{file.filename}")
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # 创建第一个版本（不存储内容，从文件路径读取）
        version = ContractVersion(
            contract_id=contract.id,
            version=1,
            file_path=file_path,
            file_type=file.content_type
        )
        db.add(version)
        db.commit()
    
    return {"id": contract.id, "title": contract.title, "status": contract.status.value}

# 获取合同详情
@router.get("/contracts/{contract_id}", response_model=dict)
async def get_contract(
    contract_id: int,
    db: Session = Depends(get_db)
):
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="合同不存在")
    
    # 获取所有版本
    versions = []
    for v in contract.versions:
        versions.append({
            "id": v.id,
            "version": v.version,
            "created_at": v.created_at,
            "file_type": v.file_type
        })
    
    return {
        "id": contract.id,
        "title": contract.title,
        "description": contract.description,
        "status": contract.status.value,
        "contract_type": contract.contract_type.value,
        "created_at": contract.created_at,
        "updated_at": contract.updated_at,
        "versions": versions
    }

# 更新合同状态
@router.patch("/contracts/{contract_id}/status", response_model=dict)
async def update_contract_status(
    contract_id: int,
    status: ContractStatus,
    db: Session = Depends(get_db)
):
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="合同不存在")
    
    contract.status = status
    db.commit()
    db.refresh(contract)
    
    return {"id": contract.id, "status": contract.status.value}

# 获取合同版本内容（文本形式）
@router.get("/contracts/{contract_id}/versions/{version_id}/content", response_model=dict)
async def get_contract_version_content(
    contract_id: int,
    version_id: int,
    db: Session = Depends(get_db)
):
    version = db.query(ContractVersion).filter(
        ContractVersion.id == version_id,
        ContractVersion.contract_id == contract_id
    ).first()
    
    if not version:
        raise HTTPException(status_code=404, detail="版本不存在")
    
    # 从文件路径读取内容
    content = ""
    if version.file_path and os.path.exists(version.file_path):
        try:
            content = extract_text_from_file(version.file_path, version.file_type)
        except Exception as e:
            print(f"读取文件内容失败: {str(e)}")
            content = "读取文件内容失败"
    
    return {
        "id": version.id,
        "version": version.version,
        "content": content,
        "file_type": version.file_type,
        "created_at": version.created_at
    }

# 获取合同版本文件流
@router.get("/contracts/{contract_id}/versions/{version_id}/file")
async def get_contract_version_file(
    contract_id: int,
    version_id: int,
    db: Session = Depends(get_db)
):
    version = db.query(ContractVersion).filter(
        ContractVersion.id == version_id,
        ContractVersion.contract_id == contract_id
    ).first()
    
    if not version:
        raise HTTPException(status_code=404, detail="版本不存在")
    
    if not version.file_path or not os.path.exists(version.file_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    
    # 打开文件并返回文件流
    return FileResponse(path=version.file_path, media_type=version.file_type, filename=os.path.basename(version.file_path))

# 更新合同版本文件
@router.put("/contracts/{contract_id}/versions/{version_id}/file", response_model=dict)
async def update_contract_version_file(
    contract_id: int,
    version_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    version = db.query(ContractVersion).filter(
        ContractVersion.id == version_id,
        ContractVersion.contract_id == contract_id
    ).first()
    
    if not version:
        raise HTTPException(status_code=404, detail="版本不存在")
    
    # 保存新文件
    file_path = os.path.join(UPLOAD_DIR, f"{contract_id}_v{version.version}_{file.filename}")
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # 更新版本信息
    version.file_path = file_path
    version.file_type = file.content_type
    db.commit()
    db.refresh(version)
    
    # 更新合同状态为已修改
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if contract:
        contract.status = ContractStatus.MODIFIED
        db.commit()
    
    return {
        "id": version.id,
        "version": version.version,
        "file_path": version.file_path,
        "file_type": version.file_type,
        "updated_at": version.created_at
    }

# 保存修改后的文本内容到文件
@router.post("/contracts/{contract_id}/versions/{version_id}/save", response_model=dict)
async def save_contract_version_content(
    contract_id: int,
    version_id: int,
    content: str = Form(...),
    db: Session = Depends(get_db)
):
    print(f"接收到保存请求:")
    print(f"  合同ID: {contract_id}")
    print(f"  版本ID: {version_id}")
    print(f"  内容长度: {len(content) if content else 0}")
    
    if not content:
        raise HTTPException(status_code=400, detail="内容不能为空")
    
    # 查找版本
    version = db.query(ContractVersion).filter(
        ContractVersion.id == version_id,
        ContractVersion.contract_id == contract_id
    ).first()
    
    if not version:
        raise HTTPException(status_code=404, detail="版本不存在")
    
    # 清理文件路径，移除多余的空格
    version.file_path = version.file_path.strip()
    print(f"保存内容到文件: {version.file_path}")
    print(f"文件类型: {version.file_type}")
    print(f"文件路径是否存在: {os.path.exists(version.file_path)}")
    print(f"文件目录是否存在: {os.path.exists(os.path.dirname(version.file_path))}")
    
    # 确保文件目录存在
    if not os.path.exists(os.path.dirname(version.file_path)):
        os.makedirs(os.path.dirname(version.file_path), exist_ok=True)
        print(f"创建文件目录: {os.path.dirname(version.file_path)}")
    
    # 保存修改后的内容到文件
    if version.file_path:
        try:
            # 确保上传目录存在
            os.makedirs(UPLOAD_DIR, exist_ok=True)
            
            # 检查文件路径是否存在
            if not os.path.exists(version.file_path):
                print(f"文件路径不存在: {version.file_path}")
                # 如果文件不存在，创建一个新文件
                print(f"创建新文件: {version.file_path}")
            
            # 根据文件类型保存内容
            if version.file_type == "text/plain" or version.file_path.endswith(".txt"):
                # 对于文本文件，直接保存
                with open(version.file_path, "w", encoding="utf-8") as f:
                    f.write(content)
                print(f"成功保存内容到文本文件: {version.file_path}")
            elif version.file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"] or version.file_path.endswith(".docx"):
                # 对于Word文件，需要使用python-docx库保存
                from docx import Document
                # 创建一个新的Word文档
                doc = Document()
                # 添加内容
                for paragraph in content.split('\n'):
                    doc.add_paragraph(paragraph)
                # 保存文档
                doc.save(version.file_path)
                print(f"成功保存内容到Word文件: {version.file_path}")
            else:
                # 对于其他文件类型，保存为文本文件
                # 首先获取文件名和扩展名
                base_name = os.path.basename(version.file_path)
                name_without_ext = os.path.splitext(base_name)[0]
                # 创建一个新的文本文件路径
                new_file_path = os.path.join(os.path.dirname(version.file_path), f"{name_without_ext}.txt")
                with open(new_file_path, "w", encoding="utf-8") as f:
                    f.write(content)
                print(f"成功保存内容到文本文件: {new_file_path}")
                # 更新版本的文件路径和文件类型
                version.file_path = new_file_path
                version.file_type = "text/plain"
                db.commit()
                print(f"更新版本文件路径为: {new_file_path}")
        except Exception as e:
            print(f"保存文件失败: {str(e)}")
            raise HTTPException(status_code=500, detail=f"保存文件失败: {str(e)}")
    else:
        # 如果没有文件路径，创建一个新的文件路径
        file_path = os.path.join(UPLOAD_DIR, f"{contract_id}_v{version.version}_modified.txt")
        try:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)
            print(f"成功创建并保存内容到文件: {file_path}")
            # 更新版本的文件路径
            version.file_path = file_path
            version.file_type = "text/plain"
            db.commit()
        except Exception as e:
            print(f"创建文件失败: {str(e)}")
            raise HTTPException(status_code=500, detail=f"创建文件失败: {str(e)}")
    
    # 更新合同状态为已修改
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if contract:
        contract.status = ContractStatus.MODIFIED
        db.commit()
    
    return {
        "id": version.id,
        "version": version.version,
        "message": "保存成功"
    }

# 上传新版本
@router.post("/contracts/{contract_id}/versions", response_model=dict)
async def upload_contract_version(
    contract_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="合同不存在")
    
    # 计算新版本号
    latest_version = contract.versions[0] if contract.versions else None
    new_version_number = (latest_version.version + 1) if latest_version else 1
    
    # 保存文件
    file_path = os.path.join(UPLOAD_DIR, f"{contract_id}_v{new_version_number}_{file.filename}")
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # 创建新版本（不存储内容，从文件路径读取）
    version = ContractVersion(
        contract_id=contract_id,
        version=new_version_number,
        file_path=file_path,
        file_type=file.content_type
    )
    db.add(version)
    db.commit()
    db.refresh(version)
    
    # 更新合同状态为已修改
    contract.status = ContractStatus.MODIFIED
    db.commit()
    
    return {
        "id": version.id,
        "version": version.version,
        "created_at": version.created_at
    }

# 版本对比
@router.get("/contracts/{contract_id}/compare", response_model=dict)
async def compare_versions(
    contract_id: int,
    version1: int,
    version2: int,
    db: Session = Depends(get_db)
):
    # 获取两个版本
    v1 = db.query(ContractVersion).filter(
        ContractVersion.contract_id == contract_id,
        ContractVersion.version == version1
    ).first()
    
    v2 = db.query(ContractVersion).filter(
        ContractVersion.contract_id == contract_id,
        ContractVersion.version == version2
    ).first()
    
    if not v1 or not v2:
        raise HTTPException(status_code=404, detail="版本不存在")
    
    # 从文件路径读取内容
    content1 = ""
    if v1.file_path and os.path.exists(v1.file_path):
        try:
            content1 = extract_text_from_file(v1.file_path, v1.file_type)
        except Exception as e:
            print(f"读取文件内容失败: {str(e)}")
            content1 = "读取文件内容失败"
    
    content2 = ""
    if v2.file_path and os.path.exists(v2.file_path):
        try:
            content2 = extract_text_from_file(v2.file_path, v2.file_type)
        except Exception as e:
            print(f"读取文件内容失败: {str(e)}")
            content2 = "读取文件内容失败"
    
    # 简化的差异对比（实际应用中需要更复杂的算法）
    diff = {
        "version1": version1,
        "version2": version2,
        "content1": content1,
        "content2": content2,
        "differences": "这里应该是详细的差异对比结果"
    }
    
    return diff

# AI交互相关路由

# 创建对话会话
@router.post("/conversations", response_model=dict)
async def create_conversation(
    session_id: str = Form(...),
    db: Session = Depends(get_db)
):
    # 检查会话是否已存在
    existing = db.query(Conversation).filter(Conversation.session_id == session_id).first()
    if existing:
        return {"session_id": existing.session_id, "created_at": existing.created_at}
    
    # 创建新会话
    conversation = Conversation(session_id=session_id)
    db.add(conversation)
    db.commit()
    db.refresh(conversation)
    
    return {"session_id": conversation.session_id, "created_at": conversation.created_at}

# 获取对话历史
@router.get("/conversations/{session_id}/messages", response_model=List[dict])
async def get_conversation_messages(
    session_id: str,
    db: Session = Depends(get_db)
):
    conversation = db.query(Conversation).filter(Conversation.session_id == session_id).first()
    if not conversation:
        raise HTTPException(status_code=404, detail="会话不存在")
    
    # 获取消息
    messages = []
    for msg in conversation.messages:
        messages.append({
            "id": msg.id,
            "role": msg.role,
            "content": msg.content,
            "created_at": msg.created_at
        })
    
    return messages

#// 发送消息（AI交互）
@router.post("/conversations/{session_id}/messages", response_model=dict)
async def send_message(
    session_id: str,
    content: str = Form(...),
    contract_id: Optional[str] = Form(None),
    contract_title: Optional[str] = Form(None),
    db: Session = Depends(get_db)
):
    # 获取或创建会话
    conversation = db.query(Conversation).filter(Conversation.session_id == session_id).first()
    if not conversation:
        conversation = Conversation(session_id=session_id)
        db.add(conversation)
        db.commit()
        db.refresh(conversation)
    
    # 保存用户消息
    user_message = Message(
        conversation_id=conversation.id,
        role="user",
        content=content
    )
    db.add(user_message)
    db.commit()
    
    # 使用DashScope API生成AI响应
    try:
        # 构建请求参数
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {settings.DASHSCOPE_API_KEY}"
        }
        
        # 获取对话历史，用于上下文
        history_messages = []
        
        # 如果有合同信息，添加合同内容到上下文
        if contract_id:
            print(f"收到合同ID: {contract_id}, 合同标题: {contract_title}")
            try:
                # 将字符串转换为整数
                contract_id_int = int(contract_id)
                # 获取合同的最新版本
                contract = db.query(Contract).filter(Contract.id == contract_id_int).first()
                
                if contract:
                    print(f"找到合同: {contract.title}")
                    if contract.versions:
                        print(f"合同版本数量: {len(contract.versions)}")
                        # 按版本号降序排序，获取最新版本
                        latest_version = max(contract.versions, key=lambda v: v.version)
                        print(f"最新版本: v{latest_version.version}")
                        
                        # 从文件路径读取内容
                        contract_content = ""
                        if latest_version.file_path and os.path.exists(latest_version.file_path):
                            try:
                                contract_content = extract_text_from_file(latest_version.file_path, latest_version.file_type)
                                print(f"成功读取合同内容，长度: {len(contract_content)} 字符")
                            except Exception as e:
                                print(f"读取文件内容失败: {str(e)}")
                                contract_content = "读取文件内容失败"
                        else:
                            print(f"文件路径不存在或为空: {latest_version.file_path}")
                            contract_content = "无法读取合同内容"
                        
                        # 将合同内容作为系统输入
                        history_messages.append({
                            "role": "system",
                            "content": f"以下是合同信息，后续的问答将基于此合同内容进行：\n\n合同标题：{contract.title}\n合同内容：{contract_content}\n\n请基于上述合同内容回答用户问题。"
                        })
                        print("已将合同内容作为系统输入")
                    else:
                        print("合同没有版本")
                else:
                    print("未找到合同")
            except ValueError:
                print(f"合同ID格式错误: {contract_id}")
            except Exception as e:
                print(f"处理合同信息时出错: {str(e)}")
        
        # 添加历史消息
        for msg in conversation.messages:
            history_messages.append({
                "role": msg.role,
                "content": msg.content
            })
        
        # 添加当前用户消息
        history_messages.append({
            "role": "user",
            "content": content
        })
        
        # 构建请求体
        data = {
            "model": settings.DASHSCOPE_MODEL,
            "messages": history_messages,
            "temperature": 0.7,
            "max_tokens": 1024
        }
        
        # 发送请求到DashScope API
        response = requests.post(
            settings.DASHSCOPE_API_URL + "/chat/completions",
            headers=headers,
            json=data
        )
        
        # 解析响应
        response_data = response.json()
        if "choices" in response_data and len(response_data["choices"]) > 0:
            ai_response = response_data["choices"][0]["message"]["content"]
        else:
            ai_response = "抱歉，我无法生成回复，请稍后再试。"
            
    except Exception as e:
        print(f"DashScope API调用失败: {str(e)}")
        ai_response = "抱歉，AI服务暂时不可用，请稍后再试。"
    
    # 保存AI消息
    ai_message = Message(
        conversation_id=conversation.id,
        role="assistant",
        content=ai_response
    )
    db.add(ai_message)
    db.commit()
    db.refresh(ai_message)
    
    return {
        "id": ai_message.id,
        "role": ai_message.role,
        "content": ai_message.content,
        "created_at": ai_message.created_at
    }

# 审核规则相关路由

# 获取审核规则
@router.get("/review-rules", response_model=List[dict])
async def get_review_rules(
    rule_type: Optional[str] = None,
    db: Session = Depends(get_db)
):
    query = db.query(ReviewRule)
    if rule_type:
        query = query.filter(ReviewRule.rule_type == rule_type)
    
    rules = query.all()
    
    result = []
    for rule in rules:
        result.append({
            "id": rule.id,
            "name": rule.name,
            "description": rule.description,
            "rule_type": rule.rule_type,
            "pattern": rule.pattern,
            "severity": rule.severity
        })
    
    return result

# 对合同进行合规性检查
@router.post("/contracts/{contract_id}/review", response_model=dict)
async def review_contract(
    contract_id: int,
    db: Session = Depends(get_db)
):
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="合同不存在")
    
    # 获取最新版本
    latest_version = contract.versions[0] if contract.versions else None
    if not latest_version:
        raise HTTPException(status_code=404, detail="合同无版本")
    
    # 从文件路径读取内容
    contract_content = ""
    if latest_version.file_path and os.path.exists(latest_version.file_path):
        try:
            contract_content = extract_text_from_file(latest_version.file_path, latest_version.file_type)
        except Exception as e:
            print(f"读取文件内容失败: {str(e)}")
            contract_content = "读取文件内容失败"
    
    # 获取所有审核规则
    rules = db.query(ReviewRule).all()
    
    # 进行检查（简化处理）
    issues = []
    for rule in rules:
        # 这里应该使用正则表达式或更复杂的规则引擎进行检查
        if rule.pattern in contract_content:
            issues.append({
                "rule_id": rule.id,
                "rule_name": rule.name,
                "description": rule.description,
                "severity": rule.severity,
                "location": "在合同内容中发现"
            })
    
    return {
        "contract_id": contract_id,
        "contract_title": contract.title,
        "issues": issues,
        "total_issues": len(issues)
    }

# 文件处理相关路由

# 上传文件
@router.post("/upload", response_model=dict)
async def upload_file(
    file: UploadFile = File(...)
):
    # 保存文件
    file_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    return {
        "filename": file.filename,
        "file_path": file_path,
        "content_type": file.content_type
    }

# 获取文件
@router.get("/files/{filename}")
async def get_file(filename: str):
    file_path = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    
    # 实际应用中应该返回文件流
    return {"file_path": file_path, "exists": True}

# 通过路径读取文件内容
@router.get("/file-content", response_model=dict)
async def get_file_content(file_path: str):
    """通过路径读取文件内容
    
    Args:
        file_path: 文件路径，相对于uploads目录
    """
    # 构建完整的文件路径
    full_path = os.path.join(UPLOAD_DIR, file_path)
    
    # 安全检查：确保文件在uploads目录内
    if not os.path.abspath(full_path).startswith(os.path.abspath(UPLOAD_DIR)):
        raise HTTPException(status_code=403, detail="无权访问该文件路径")
    
    # 检查文件是否存在
    if not os.path.exists(full_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    
    # 检查是否是文件
    if not os.path.isfile(full_path):
        raise HTTPException(status_code=400, detail="路径不是文件")
    
    # 提取文件类型
    import mimetypes
    file_type, _ = mimetypes.guess_type(full_path)
    if file_type is None:
        # 根据文件扩展名猜测类型
        if full_path.endswith('.pdf'):
            file_type = 'application/pdf'
        elif full_path.endswith('.docx'):
            file_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif full_path.endswith('.doc'):
            file_type = 'application/msword'
        elif full_path.endswith('.txt'):
            file_type = 'text/plain'
        else:
            file_type = 'application/octet-stream'
    
    # 提取文件内容
    try:
        content = extract_text_from_file(full_path, file_type)
    except Exception as e:
        print(f"提取文件内容失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"提取文件内容失败: {str(e)}")
    
    return {
        "file_path": file_path,
        "full_path": full_path,
        "file_type": file_type,
        "content": content
    }